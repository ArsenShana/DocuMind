"""Generate project report in Word format."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_cell_shading(cell, color):
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h


def add_table_row(table, cells, header=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        for p in cell.paragraphs:
            p.style.font.size = Pt(10)
        if header:
            set_cell_shading(cell, "1a1a2e")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== TITLE PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DocuMind AI")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x06, 0xb6, 0xd4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Intelligent Document Analysis & Validation System")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x64, 0x64, 0x80)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Technical Project Report")
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. Problem Statement",
        "3. System Architecture",
        "4. Technology Stack",
        "5. Core Components",
        "   5.1 Document Parser",
        "   5.2 LLM Classifier",
        "   5.3 Data Extractor",
        "   5.4 Cross-Document Validator",
        "   5.5 AI Summarizer",
        "6. API Design",
        "7. Frontend Application",
        "8. Data Models & Schemas",
        "9. Deployment",
        "10. Testing & Results",
        "11. Future Improvements",
        "12. Conclusion",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ===== 1. INTRODUCTION =====
    add_heading_styled(doc, "1. Introduction", level=1)
    doc.add_paragraph(
        "DocuMind AI is a backend service powered by Large Language Models (LLM) designed for "
        "automatic processing and analysis of documents such as invoices, receipts, and contracts. "
        "The system classifies document types, extracts structured data with confidence scores, "
        "performs cross-document validation to detect logical and financial discrepancies, and "
        "generates AI-powered analytical summaries."
    )
    doc.add_paragraph(
        "The solution employs a hybrid approach combining the capabilities of large language models "
        "(OpenAI GPT-4o-mini) with deterministic rule-based logic. This enables effective processing "
        "of unstructured data from real-world documents while ensuring reliable, interpretable results."
    )
    doc.add_paragraph(
        "The project demonstrates practical application of LLM for document understanding tasks, "
        "as well as skills in designing robust backend services for working with unstructured data."
    )

    # ===== 2. PROBLEM STATEMENT =====
    add_heading_styled(doc, "2. Problem Statement", level=1)
    doc.add_paragraph(
        "Organizations process thousands of documents daily — invoices, receipts, contracts. "
        "Manual review is slow, error-prone, and expensive. Key challenges include:"
    )
    problems = [
        "Document classification requires manual sorting by type",
        "Data extraction from unstructured text is labor-intensive",
        "Cross-referencing documents for discrepancies is often missed",
        "Scanned documents and photos require OCR before processing",
        "No standardized output format for downstream systems",
    ]
    for p in problems:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_paragraph(
        "DocuMind AI addresses all these challenges by providing an automated, AI-driven pipeline "
        "that processes documents end-to-end with minimal human intervention."
    )

    # ===== 3. ARCHITECTURE =====
    add_heading_styled(doc, "3. System Architecture", level=1)
    doc.add_paragraph(
        "The system follows a modular pipeline architecture where each stage processes the document "
        "sequentially, with results passed to the next stage:"
    )

    arch = doc.add_paragraph()
    arch.style.font.name = 'Courier New'
    arch.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = arch.add_run(
        "Upload (PDF/Image)\n"
        "      |\n"
        "      v\n"
        "[Document Parser] -- PyMuPDF + Tesseract OCR\n"
        "      |\n"
        "      v\n"
        "[LLM Classifier] -- GPT-4o-mini -> invoice/receipt/contract\n"
        "      |\n"
        "      v\n"
        "[Data Extractor] -- GPT-4o-mini -> structured JSON\n"
        "      |\n"
        "      v\n"
        "[Validator] -- Rule Engine + LLM -> discrepancies\n"
        "      |\n"
        "      v\n"
        "[AI Summarizer] -- GPT-4o-mini -> summary + risk\n"
        "      |\n"
        "      v\n"
        "Structured JSON Response"
    )
    run.font.size = Pt(9)

    doc.add_paragraph(
        "Each component is implemented as an independent service, enabling easy testing, "
        "replacement, and scaling of individual pipeline stages."
    )

    # ===== 4. TECH STACK =====
    add_heading_styled(doc, "4. Technology Stack", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(["Component", "Technology", "Purpose"]):
        hdr[i].text = text
        set_cell_shading(hdr[i], "1a1a2e")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True

    stack = [
        ("Backend Framework", "FastAPI (async)", "REST API with auto-generated docs"),
        ("LLM Integration", "OpenAI GPT-4o-mini", "Classification, extraction, summarization"),
        ("Schema Validation", "Pydantic v2", "20+ strict data models"),
        ("Document Parsing", "PyMuPDF", "PDF text extraction"),
        ("OCR Engine", "Tesseract OCR", "Image & scanned PDF recognition"),
        ("Image Processing", "Pillow (PIL)", "Image preprocessing for OCR"),
        ("HTTP Client", "httpx", "Async OpenAI API calls"),
        ("Frontend", "React + Vite", "Landing page + Dashboard"),
        ("Animations", "Framer Motion", "UI animations and transitions"),
        ("Containerization", "Docker + Compose", "Production deployment"),
        ("Web Server", "Nginx", "Reverse proxy + static files"),
    ]
    for row_data in stack:
        row = table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    # ===== 5. CORE COMPONENTS =====
    add_heading_styled(doc, "5. Core Components", level=1)

    # 5.1
    add_heading_styled(doc, "5.1 Document Parser", level=2)
    doc.add_paragraph(
        "The document parser is the entry point of the pipeline. It handles multiple input formats "
        "and ensures text is extracted regardless of the document source:"
    )
    parser_features = [
        "PDF with embedded text: Direct extraction via PyMuPDF (fast, accurate)",
        "Scanned PDF (no text layer): Automatic fallback to Tesseract OCR at 300 DPI",
        "Photos (PNG, JPG, TIFF, BMP, WebP): Full OCR with auto-rotation (EXIF)",
        "Multi-language support: English and Russian recognition",
        "Metadata extraction: Page count, file size, image dimensions",
    ]
    for f in parser_features:
        doc.add_paragraph(f, style='List Bullet')

    # 5.2
    add_heading_styled(doc, "5.2 LLM Classifier", level=2)
    doc.add_paragraph(
        "The classifier sends extracted text to GPT-4o-mini with a structured prompt. "
        "The LLM returns a JSON response with:"
    )
    class_fields = [
        "document_type — one of: invoice, receipt, contract, unknown",
        "confidence — float 0.0 to 1.0 indicating classification certainty",
        "reasoning — explanation of why this classification was chosen",
    ]
    for f in class_fields:
        doc.add_paragraph(f, style='List Bullet')
    doc.add_paragraph(
        "The temperature is set to 0.1 for deterministic output, and response_format is enforced "
        "as JSON to prevent hallucinated text outside the schema."
    )

    # 5.3
    add_heading_styled(doc, "5.3 Data Extractor", level=2)
    doc.add_paragraph(
        "Based on the classification result, a type-specific extraction prompt is selected. "
        "Each document type has its own Pydantic model:"
    )
    extract_models = [
        "InvoiceData — invoice_number, dates, vendor/buyer, line_items, subtotal, tax, total",
        "ReceiptData — store info, items, amounts, payment method, transaction ID",
        "ContractData — parties, dates, total value, key terms, obligations, termination clause",
    ]
    for m in extract_models:
        doc.add_paragraph(m, style='List Bullet')
    doc.add_paragraph(
        "Every monetary value includes a confidence score. The overall extraction confidence "
        "determines the confidence level: HIGH (>=0.8), MEDIUM (>=0.5), or LOW (<0.5)."
    )

    # 5.4
    add_heading_styled(doc, "5.4 Cross-Document Validator", level=2)
    doc.add_paragraph(
        "The validation engine uses a hybrid approach — the key innovation of the system:"
    )

    doc.add_paragraph("Rule-Based Checks (Deterministic):", style='List Bullet')
    rules = [
        "Amount comparison between documents (flag >10% difference as critical)",
        "Date consistency verification",
        "Company/party name matching (case-insensitive)",
        "Internal math validation: subtotal + tax = total",
    ]
    for r in rules:
        p = doc.add_paragraph(r)
        p.paragraph_format.left_indent = Cm(2)

    doc.add_paragraph("LLM-Based Analysis (Semantic):", style='List Bullet')
    llm_checks = [
        "Complex logical inconsistencies that rules can't catch",
        "Duplicate document detection",
        "Missing required fields identification",
        "Overall risk assessment and narrative summary",
    ]
    for r in llm_checks:
        p = doc.add_paragraph(r)
        p.paragraph_format.left_indent = Cm(2)

    doc.add_paragraph(
        "Results are merged with deduplication (by field + type), and a composite risk score "
        "is calculated: risk = critical_count * 0.3 + warning_count * 0.1 + llm_risk * 0.3, "
        "capped at 1.0."
    )

    # 5.5
    add_heading_styled(doc, "5.5 AI Summarizer", level=2)
    doc.add_paragraph(
        "Generates a concise summary of each document including: "
        "2-3 sentence overview, key findings list, issues found list, and risk level "
        "(info/warning/critical). Uses a dedicated LLM prompt with temperature 0.2 "
        "for balanced yet consistent output."
    )

    # ===== 6. API DESIGN =====
    add_heading_styled(doc, "6. API Design", level=1)
    doc.add_paragraph(
        "The REST API is built with FastAPI and provides automatic OpenAPI/Swagger documentation. "
        "All responses are strictly typed with Pydantic models."
    )

    api_table = doc.add_table(rows=1, cols=3)
    api_table.style = 'Table Grid'
    hdr = api_table.rows[0].cells
    for i, text in enumerate(["Method", "Endpoint", "Description"]):
        hdr[i].text = text
        set_cell_shading(hdr[i], "1a1a2e")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True

    endpoints = [
        ("POST", "/api/documents/upload", "Upload document (PDF, PNG, JPG)"),
        ("POST", "/api/documents/analyze/{id}", "Full pipeline: classify + extract + summarize"),
        ("POST", "/api/documents/classify/{id}", "Classify document type only"),
        ("POST", "/api/documents/extract/{id}", "Extract structured data only"),
        ("POST", "/api/documents/validate", "Cross-validate multiple documents"),
        ("GET", "/api/documents/", "List all uploaded documents"),
        ("GET", "/api/documents/{id}", "Get document details and results"),
        ("GET", "/health", "Service health check"),
    ]
    for row_data in endpoints:
        row = api_table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = text

    # ===== 7. FRONTEND =====
    add_heading_styled(doc, "7. Frontend Application", level=1)
    doc.add_paragraph(
        "The frontend is a React single-page application built with Vite, "
        "featuring Framer Motion animations and a dark theme with cyan/violet accents."
    )

    add_heading_styled(doc, "Landing Page", level=3)
    landing = [
        "Hero section with animated gradient orbs and call-to-action",
        "6 feature cards with hover animations and gradient borders",
        "How It Works — 4-step visual guide with connected dots",
        "Pipeline visualization showing the full processing chain",
        "CTA section with glow effects",
        "Fully responsive design (mobile, tablet, desktop)",
    ]
    for item in landing:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_styled(doc, "Dashboard", level=3)
    dash = [
        "Upload & Analyze — drag-and-drop file upload with real-time analysis progress, "
        "results displayed as classification card with confidence bar, "
        "extracted data grid with line items table, AI summary with key findings and issues",
        "Cross-Validate — select multiple analyzed documents, run hybrid validation, "
        "view discrepancy cards with severity levels, risk meter, and suggestions",
        "Documents — list of all uploaded documents with status (analyzed/pending)",
    ]
    for item in dash:
        doc.add_paragraph(item, style='List Bullet')

    # ===== 8. DATA MODELS =====
    add_heading_styled(doc, "8. Data Models & Schemas", level=1)
    doc.add_paragraph(
        "The system uses 20+ Pydantic models for strict data validation. Key schemas:"
    )
    schemas = [
        ("ClassificationResult", "document_type, confidence, reasoning"),
        ("ExtractionResult", "document_type, raw_text, extracted_data, entities, confidence"),
        ("InvoiceData", "invoice_number, dates, vendor/buyer, line_items, amounts"),
        ("ReceiptData", "store, date/time, items, amounts, payment_method"),
        ("ContractData", "parties, dates, value, terms, obligations"),
        ("ValidationResult", "is_valid, discrepancies[], summary, risk_score"),
        ("Discrepancy", "type, severity, description, values, suggestion"),
        ("DocumentSummary", "summary, key_findings[], issues_found[], risk_level"),
        ("MoneyAmount", "value, currency, confidence"),
        ("FullAnalysisResult", "classification + extraction + summary combined"),
    ]

    schema_table = doc.add_table(rows=1, cols=2)
    schema_table.style = 'Table Grid'
    hdr = schema_table.rows[0].cells
    for i, text in enumerate(["Schema", "Key Fields"]):
        hdr[i].text = text
        set_cell_shading(hdr[i], "1a1a2e")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
    for name, fields in schemas:
        row = schema_table.add_row()
        row.cells[0].text = name
        row.cells[1].text = fields

    # ===== 9. DEPLOYMENT =====
    add_heading_styled(doc, "9. Deployment", level=1)
    doc.add_paragraph(
        "The application is containerized using Docker Compose with two services:"
    )
    deploy = [
        "documind (backend) — Python 3.12, FastAPI, Uvicorn, PyMuPDF, Tesseract OCR",
        "frontend — Node.js build → Nginx serving static files + reverse proxy to API",
    ]
    for d in deploy:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph(
        "Configuration is managed via environment variables (.env file). "
        "The only external dependency is the OpenAI API key. "
        "Uploaded files are stored in a Docker volume (./uploads)."
    )

    # ===== 10. TESTING =====
    add_heading_styled(doc, "10. Testing & Results", level=1)
    doc.add_paragraph("The system was tested with sample documents:")

    add_heading_styled(doc, "Classification Test", level=3)
    doc.add_paragraph(
        "Invoice PDF uploaded → classified as 'invoice' with confidence 1.0. "
        "The reasoning correctly identified invoice number, vendor info, line items, and totals."
    )

    add_heading_styled(doc, "Extraction Test", level=3)
    doc.add_paragraph(
        "All fields extracted correctly: invoice number (INV-2026-0042), dates, "
        "vendor/buyer names and addresses, 3 line items with quantities and amounts, "
        "subtotal ($4,500.00), tax ($382.50), total ($4,882.50). Overall confidence: 0.90."
    )

    add_heading_styled(doc, "OCR Test", level=3)
    doc.add_paragraph(
        "PNG image of an invoice uploaded → Tesseract OCR recognized text → "
        "classified as invoice (confidence 0.9) → extracted data correctly including "
        "company name and total amount."
    )

    add_heading_styled(doc, "Cross-Validation Test", level=3)
    doc.add_paragraph(
        "Invoice and mismatched receipt cross-validated. System detected:"
    )
    findings = [
        "CRITICAL: Tax amount mismatch ($382.50 vs $405.00)",
        "CRITICAL: Total amount mismatch ($4,882.50 vs $4,905.00)",
        "WARNING: Date mismatch (2026-03-15 vs 2026-03-16)",
        "INFO: Company name format difference (TechCorp Solutions LLC vs TECHCORP SOLUTIONS)",
        "Risk score: 1.0 (high risk)",
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Bullet')

    # ===== 11. FUTURE =====
    add_heading_styled(doc, "11. Future Improvements", level=1)
    improvements = [
        "Database storage (PostgreSQL) for persistent document management",
        "FAISS vector search for semantic document similarity",
        "User authentication and multi-tenant support",
        "Batch processing for bulk document uploads",
        "Custom document type training (purchase orders, tax forms)",
        "Webhook notifications for async processing results",
        "PDF annotation — highlight extracted fields in original document",
        "Multi-language OCR expansion (Chinese, Arabic, etc.)",
    ]
    for item in improvements:
        doc.add_paragraph(item, style='List Bullet')

    # ===== 12. CONCLUSION =====
    add_heading_styled(doc, "12. Conclusion", level=1)
    doc.add_paragraph(
        "DocuMind AI demonstrates a practical, production-ready approach to intelligent document "
        "processing. By combining the semantic understanding capabilities of large language models "
        "with deterministic rule-based validation, the system achieves both flexibility and reliability."
    )
    doc.add_paragraph(
        "The modular architecture ensures each component can be independently tested, upgraded, "
        "or replaced. The strict Pydantic schema validation guarantees consistent, machine-readable "
        "output suitable for integration with enterprise systems."
    )
    doc.add_paragraph(
        "The hybrid validation engine is the key differentiator — while LLMs excel at semantic "
        "understanding, rule-based checks provide guaranteed detection of mathematical and logical "
        "errors. Together, they deliver comprehensive document intelligence with measurable "
        "confidence scoring."
    )

    # Save
    doc.save("/root/DocuMindAI/DocuMind_AI_Report.docx")
    print("Report saved: DocuMind_AI_Report.docx")


if __name__ == "__main__":
    main()
